"""
DocumentProcessor adapter implementing DocumentProcessorPort.

Handles text extraction from multiple formats and intelligent chunking.
"""

import re
from pathlib import Path
from typing import Optional

# PDF extraction
try:
    from pypdf import PdfReader
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

# DOCX extraction
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from memoria.domain.entities import Chunk, Document
from memoria.domain.ports.document_processor import DocumentProcessorPort


class DocumentProcessorAdapter:
    """
    Adapter for document processing operations.

    Implements DocumentProcessorPort protocol with support for:
    - Text extraction from PDF, DOCX, TXT, MD
    - Smart chunking with word boundaries
    - Metadata extraction
    """

    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
    ) -> None:
        """
        Initialize document processor adapter.

        Args:
            chunk_size: Default chunk size in characters
            chunk_overlap: Default overlap between chunks
        """
        self._chunk_size = chunk_size
        self._chunk_overlap = chunk_overlap

    def extract_text(self, file_path: Path) -> str:
        """
        Extract text from a file.

        Args:
            file_path: Path to the file

        Returns:
            Extracted text content

        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file doesn't exist
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        suffix = file_path.suffix.lower()

        if suffix == ".txt":
            return self._extract_txt(file_path)
        elif suffix == ".md":
            return self._extract_md(file_path)
        elif suffix == ".pdf":
            return self._extract_pdf(file_path)
        elif suffix == ".docx":
            return self._extract_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")

    def _extract_txt(self, file_path: Path) -> str:
        """Extract text from plain text file."""
        return file_path.read_text(encoding="utf-8", errors="ignore")

    def _extract_md(self, file_path: Path) -> str:
        """Extract text from markdown file."""
        # For now, treat markdown as plain text
        # Could enhance with markdown parsing later
        return file_path.read_text(encoding="utf-8", errors="ignore")

    def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        if not HAS_PDF:
            raise ImportError("pypdf not installed. Install with: pip install pypdf")

        try:
            reader = PdfReader(str(file_path))
            text_parts = []

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            return "\n\n".join(text_parts)

        except Exception as e:
            raise ValueError(f"Failed to extract PDF text: {e}")

    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        if not HAS_DOCX:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")

        try:
            doc = DocxDocument(str(file_path))
            text_parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
            return "\n\n".join(text_parts)

        except Exception as e:
            raise ValueError(f"Failed to extract DOCX text: {e}")

    def chunk_text(
        self,
        text: str,
        chunk_size: Optional[int] = None,
        overlap: Optional[int] = None,
    ) -> list[Chunk]:
        """
        Split text into overlapping chunks with word-boundary awareness.

        Attempts to break chunks at word boundaries (spaces) when possible
        to avoid splitting words in the middle.

        Args:
            text: Text to chunk
            chunk_size: Size of each chunk (uses default if None)
            overlap: Overlap between chunks (uses default if None)

        Returns:
            List of Chunk objects
        """
        chunk_size = chunk_size if chunk_size is not None else self._chunk_size
        overlap = overlap if overlap is not None else self._chunk_overlap

        if not text:
            return []

        chunks: list[Chunk] = []
        start = 0

        while start < len(text):
            # Calculate end position for this chunk
            end = min(start + chunk_size, len(text))

            # Try to break at word boundary if not at end of text
            if end < len(text) and not text[end].isspace():
                # Look backwards for a space
                space_pos = text.rfind(" ", start, end)
                if space_pos > start:  # Found a space
                    end = space_pos

            # Extract chunk text
            chunk_text = text[start:end]

            if chunk_text:
                chunk = Chunk(
                    text=chunk_text,
                    start_pos=start,
                    end_pos=end,
                    metadata={},
                )
                chunks.append(chunk)

            # Calculate next start position with overlap
            next_start = end - overlap

            # Ensure forward progress (avoid infinite loop)
            if next_start <= start:
                next_start = start + 1

            start = next_start

            # If we've reached or passed the end, stop
            if start >= len(text):
                break

        return chunks

    def process_document(
        self,
        file_path: Path,
        chunk_size: Optional[int] = None,
        overlap: Optional[int] = None,
    ) -> list[Document]:
        """
        Process a document: extract text and chunk it.

        Args:
            file_path: Path to document
            chunk_size: Chunk size (uses default if None)
            overlap: Overlap size (uses default if None)

        Returns:
            List of Document objects (one per chunk)
        """
        # Extract text
        text = self.extract_text(file_path)

        # Chunk text
        chunks = self.chunk_text(text, chunk_size, overlap)

        # Create documents
        documents = []
        for i, chunk in enumerate(chunks):
            doc = Document(
                id=f"{file_path.name}_{i}",
                content=chunk.text,
                metadata={
                    "source": file_path.name,
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "file_path": str(file_path),
                },
                embedding=None,  # Embeddings generated separately
            )
            documents.append(doc)

        return documents

    def supported_formats(self) -> list[str]:
        """
        Get list of supported file formats.

        Returns:
            List of file extensions (with dots)
        """
        formats = [".txt", ".md"]

        if HAS_PDF:
            formats.append(".pdf")

        if HAS_DOCX:
            formats.append(".docx")

        return formats
